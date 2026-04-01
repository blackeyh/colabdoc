"""
Integrate AI documentation into backend_documentation.docx.
Uses lxml OxmlElement to insert tables without calling doc.add_table().
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
import re

DOC_PATH = "backend_documentation.docx"
doc = Document(DOC_PATH)
body = doc.element.body


# ─── Helpers ────────────────────────────────────────────────────────────────

def make_paragraph_xml(text, style_name=None, bold=False):
    """Return a <w:p> OxmlElement with text."""
    p = OxmlElement("w:p")
    if style_name:
        pPr = OxmlElement("w:pPr")
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qn("w:val"), style_name)
        pPr.append(pStyle)
        p.append(pPr)
    r = OxmlElement("w:r")
    if bold:
        rPr = OxmlElement("w:rPr")
        b = OxmlElement("w:b")
        rPr.append(b)
        r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    if text.startswith(" ") or text.endswith(" "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    p.append(r)
    return p


def make_table_xml(rows_data, col_widths=None):
    """
    Build a <w:tbl> from rows_data: list of rows, each a list of cell texts.
    col_widths: list of ints in twentieths-of-a-point (twips). Optional.
    """
    tbl = OxmlElement("w:tbl")
    tblPr = OxmlElement("w:tblPr")
    tblStyle = OxmlElement("w:tblStyle")
    tblStyle.set(qn("w:val"), "TableGrid")
    tblPr.append(tblStyle)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "0")
    tblW.set(qn("w:type"), "auto")
    tblPr.append(tblW)
    tbl.append(tblPr)

    tblGrid = OxmlElement("w:tblGrid")
    ncols = max(len(row) for row in rows_data)
    for ci in range(ncols):
        gridCol = OxmlElement("w:gridCol")
        if col_widths and ci < len(col_widths):
            gridCol.set(qn("w:w"), str(col_widths[ci]))
        tblGrid.append(gridCol)
    tbl.append(tblGrid)

    for ri, row_cells in enumerate(rows_data):
        tr = OxmlElement("w:tr")
        for ci, cell_text in enumerate(row_cells):
            tc = OxmlElement("w:tc")
            if col_widths and ci < len(col_widths):
                tcPr = OxmlElement("w:tcPr")
                tcW = OxmlElement("w:tcW")
                tcW.set(qn("w:w"), str(col_widths[ci]))
                tcW.set(qn("w:type"), "dxa")
                tcPr.append(tcW)
                tc.append(tcPr)
            p = OxmlElement("w:p")
            r = OxmlElement("w:r")
            t = OxmlElement("w:t")
            t.text = cell_text
            if cell_text.startswith(" ") or cell_text.endswith(" "):
                t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            r.append(t)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)

    return tbl


def insert_after(ref_el, *new_els):
    """Insert new_els into body immediately after ref_el."""
    parent = ref_el.getparent()
    idx = list(parent).index(ref_el)
    for offset, el in enumerate(new_els):
        parent.insert(idx + 1 + offset, el)


def set_cell_text(tc, text):
    """Replace all text in a table cell with new text."""
    # Remove existing paragraphs
    for p in tc.findall(qn("w:p")):
        tc.remove(p)
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    p.append(r)
    tc.append(p)


# ─── Get current body children list ─────────────────────────────────────────
children = list(body)


# ═══════════════════════════════════════════════════════════════════════════
# STEP 1 — Remove the appended "Section 2.2 (continued)" block (elements 287-351)
# ═══════════════════════════════════════════════════════════════════════════
# The appended block starts with an empty paragraph at index 286, then the
# heading "Section 2.2 (continued)" at 287, through to index 351 (File Locations).
# The sectPr at 352 must be kept.

print("Step 1: Removing appended section...")
# Find the "Section 2.2 (continued)" paragraph
to_remove = []
found = False
for child in list(body):
    if not found:
        tag = child.tag.split("}")[1]
        if tag == "p":
            text = "".join(r.text or "" for r in child.findall(".//" + qn("w:t")))
            if "Section 2.2 (continued)" in text:
                found = True
                to_remove.append(child)
    else:
        tag = child.tag.split("}")[1]
        if tag == "sectPr":
            break
        to_remove.append(child)

# Also remove the empty paragraph just before the heading
if to_remove:
    first = to_remove[0]
    prev = first.getprevious()
    if prev is not None:
        ptag = prev.tag.split("}")[1]
        if ptag == "p":
            ptext = "".join(r.text or "" for r in prev.findall(".//" + qn("w:t")))
            if not ptext.strip():
                to_remove.insert(0, prev)

for el in to_remove:
    body.remove(el)

print(f"  Removed {len(to_remove)} elements")

# Refresh children
children = list(body)


# ═══════════════════════════════════════════════════════════════════════════
# STEP 2 — Add "7. AI Module" paragraph + 4×2 module table after Table 5
# ═══════════════════════════════════════════════════════════════════════════
print("Step 2: Adding AI Module feature decomposition entry...")

# Find Table 5 (Database Module - 4 rows x 2 cols) by locating the "6. Database Module" paragraph
db_module_el = None
for child in list(body):
    tag = child.tag.split("}")[1]
    if tag == "p":
        text = "".join(r.text or "" for r in child.findall(".//" + qn("w:t")))
        if "6. Database Module" in text:
            db_module_el = child
            break

if db_module_el is None:
    print("  WARNING: Could not find '6. Database Module' paragraph!")
else:
    # The table immediately follows the paragraph
    db_table = db_module_el.getnext()
    # After the table is an empty paragraph
    empty_after_db = db_table.getnext()

    # Build paragraph "7. AI Module"
    p_ai_module = make_paragraph_xml("7. AI Module")

    # Build 4x2 module table
    ai_module_rows = [
        ["Field", "Details"],
        ["What it does",
         "Accepts AI assist requests from authenticated users, builds prompts based on the "
         "selected action, calls the locally-hosted LM Studio LLM via the OpenAI-compatible "
         "API, persists every interaction to the database, and returns the generated suggestion "
         "to the client."],
        ["What it depends on",
         "Auth module (JWT validation), Documents module (permission check), "
         "Database module (SQLAlchemy session), LM Studio (local LLM server at "
         "http://127.0.0.1:1234/v1), openai Python package."],
        ["What it exposes",
         "POST /documents/{doc_id}/ai/assist — request AI assistance on selected text.\n"
         "GET /documents/{doc_id}/ai/history — retrieve the 20 most recent AI interactions "
         "the current user made on the document."],
    ]
    tbl_ai_module = make_table_xml(ai_module_rows, col_widths=[1800, 6300])

    # Empty paragraph separator
    p_empty = make_paragraph_xml("")

    # Insert: after the empty paragraph that follows the Database Module table
    # Order: ... db_table, empty_after_db, [p_empty, p_ai_module, tbl_ai_module, p_empty2]
    insert_after(empty_after_db, p_empty, p_ai_module, tbl_ai_module, make_paragraph_xml(""))
    print("  Done.")


# ═══════════════════════════════════════════════════════════════════════════
# STEP 3 — Append AI endpoints to Table 6 (Endpoint Summary cell)
# ═══════════════════════════════════════════════════════════════════════════
print("Step 3: Updating Endpoint Summary table...")

# Find Table 6 by locating "Endpoint Summary" paragraph
endpoint_summary_tbl = None
for child in list(body):
    tag = child.tag.split("}")[1]
    if tag == "p":
        text = "".join(r.text or "" for r in child.findall(".//" + qn("w:t")))
        if text.strip() == "Endpoint Summary":
            # Next sibling should be the table
            endpoint_summary_tbl = child.getnext()
            break

if endpoint_summary_tbl is None:
    print("  WARNING: Could not find Endpoint Summary table!")
else:
    tc = endpoint_summary_tbl.find(".//" + qn("w:tc"))
    # Get existing paragraphs
    existing_paras = tc.findall(qn("w:p"))
    # Add a new paragraph for AI endpoints
    for p in existing_paras:
        t_els = p.findall(".//" + qn("w:t"))
        last_text = "".join(t.text or "" for t in t_els)

    # Append AI section
    p_new = OxmlElement("w:p")
    r_new = OxmlElement("w:r")
    t_new = OxmlElement("w:t")
    t_new.text = ("AI  "
                  "POST    /documents/{id}/ai/assist  "
                  "GET     /documents/{id}/ai/history")
    t_new.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r_new.append(t_new)
    p_new.append(r_new)
    tc.append(p_new)
    print("  Done.")


# ═══════════════════════════════════════════════════════════════════════════
# STEP 4 — Add 2 API contract tables after the last versions endpoint table
# ═══════════════════════════════════════════════════════════════════════════
print("Step 4: Adding AI endpoint API contract tables...")

# Find the "POST /documents/{id}/versions/restore/{version_number}" paragraph
restore_para = None
for child in list(body):
    tag = child.tag.split("}")[1]
    if tag == "p":
        text = "".join(r.text or "" for r in child.findall(".//" + qn("w:t")))
        if "versions/restore" in text:
            restore_para = child
            break

if restore_para is None:
    print("  WARNING: Could not find restore versions paragraph!")
else:
    restore_tbl = restore_para.getnext()
    empty_after_restore = restore_tbl.getnext()

    assist_text = (
        "Receives:\n"
        '  { "selected_text": "The text the user highlighted.", "action": "rewrite", "context": "optional surrounding text" }\n'
        "\n"
        "Returns:\n"
        '  { "id": 42, "action": "rewrite", "suggestion": "Improved version of the text...", "status": "completed" }\n'
        "\n"
        "Errors:\n"
        "  400 → selected_text is empty or action is invalid  |  "
        "403 → no permission on document  |  "
        "404 → document not found  |  "
        "502 → LM Studio unreachable or returned error"
    )

    history_text = (
        "Receives:\n"
        "  nothing (JWT token in header)\n"
        "\n"
        "Returns:\n"
        '  { "history": [ { "id": 42, "action": "summarize", "selected_text": "first 100 chars...", '
        '"suggestion": "...", "status": "completed", "created_at": "2026-03-01T10:00:00Z" } ] }\n'
        "\n"
        "Errors:\n"
        "  403 → no permission on document  |  404 → document not found"
    )

    p_assist_heading = make_paragraph_xml("POST /documents/{id}/ai/assist")
    tbl_assist = make_table_xml([[assist_text]])
    p_empty1 = make_paragraph_xml("")

    p_history_heading = make_paragraph_xml("GET /documents/{id}/ai/history")
    tbl_history = make_table_xml([[history_text]])
    p_empty2 = make_paragraph_xml("")

    insert_after(
        empty_after_restore,
        p_empty1,
        p_assist_heading,
        tbl_assist,
        make_paragraph_xml(""),
        p_history_heading,
        tbl_history,
        p_empty2,
    )
    print("  Done.")


# ═══════════════════════════════════════════════════════════════════════════
# STEP 5 — Update Table 30 (AIInteraction schema)
# ═══════════════════════════════════════════════════════════════════════════
print("Step 5: Updating AIInteraction schema table...")

ai_interaction_tbl = None
for child in list(body):
    tag = child.tag.split("}")[1]
    if tag == "p":
        text = "".join(r.text or "" for r in child.findall(".//" + qn("w:t")))
        if text.strip() == "AI Interaction":
            ai_interaction_tbl = child.getnext()
            break

if ai_interaction_tbl is None:
    print("  WARNING: Could not find AI Interaction table!")
else:
    tc = ai_interaction_tbl.find(".//" + qn("w:tc"))
    new_text = (
        "AIInteraction\n"
        "────────────────────────────────────────\n"
        "id             INT              Primary key\n"
        "document_id    INT              Foreign key → Document.id\n"
        "user_id        INT              Foreign key → User.id\n"
        "selected_text  TEXT             The text the user highlighted (nullable)\n"
        "action         VARCHAR          'rewrite' | 'summarize' | 'translate' | 'restructure' (nullable)\n"
        "suggestion     TEXT             AI-generated response (nullable)\n"
        "status         VARCHAR          'pending' | 'completed' | 'failed' (nullable)\n"
        "created_at     TIMESTAMPTZ      Time of invocation (UTC, non-null)"
    )
    set_cell_text(tc, new_text)
    print("  Done.")


# ═══════════════════════════════════════════════════════════════════════════
# STEP 6 — Update Privacy Considerations bullet 1
# ═══════════════════════════════════════════════════════════════════════════
print("Step 6: Updating privacy considerations bullet...")

for child in list(body):
    tag = child.tag.split("}")[1]
    if tag == "p":
        text = "".join(r.text or "" for r in child.findall(".//" + qn("w:t")))
        if "Document content sent to the AI Service" in text:
            # Replace the text run(s) in this paragraph
            for t_el in child.findall(".//" + qn("w:t")):
                t_el.getparent().remove(t_el) if t_el.getparent().tag.split("}")[1] == "r" else None
            # Clear all runs and set new text
            for r_el in child.findall(qn("w:r")):
                child.remove(r_el)
            new_r = OxmlElement("w:r")
            new_t = OxmlElement("w:t")
            new_t.text = (
                "Document content is sent to the locally-hosted LM Studio LLM for processing. "
                "Since the model runs on-premises, no user content leaves the local network. "
                "Users should be informed that their selected text is processed by the AI assistant."
            )
            new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            new_r.append(new_t)
            child.append(new_r)
            print("  Done.")
            break


# ═══════════════════════════════════════════════════════════════════════════
# STEP 7 — Update Configuration Management paragraph
# ═══════════════════════════════════════════════════════════════════════════
print("Step 7: Updating Configuration Management paragraph...")

for child in list(body):
    tag = child.tag.split("}")[1]
    if tag == "p":
        text = "".join(r.text or "" for r in child.findall(".//" + qn("w:t")))
        if "All secrets and credentials are stored in a .env file" in text:
            for r_el in child.findall(qn("w:r")):
                child.remove(r_el)
            new_r = OxmlElement("w:r")
            new_t = OxmlElement("w:t")
            new_t.text = (
                "All secrets and credentials are stored in a .env file that is never committed to "
                "the repository. A .env.example file documents which variables are required without "
                "exposing real values. This applies to database credentials, JWT secrets, and LM Studio "
                "configuration. The LM Studio variables are: LM_STUDIO_BASE_URL (default: "
                "http://127.0.0.1:1234/v1) and LM_STUDIO_MODEL (set to the model identifier shown in "
                "LM Studio's loaded model list)."
            )
            new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            new_r.append(new_t)
            child.append(new_r)
            print("  Done.")
            break


# ═══════════════════════════════════════════════════════════════════════════
# STEP 8 — Add ADR 5 after the ADR 4 table and its content
# ═══════════════════════════════════════════════════════════════════════════
print("Step 8: Adding ADR 5 (LM Studio)...")

# Find "ADR 4 — Store Full Document Copies" paragraph to locate the ADR 4 summary table
# Then find the last paragraph of ADR 4 content (Alternatives Considered section)
# We insert ADR 5 after the last paragraph of ADR 4

# Strategy: find the ADR 4 summary table, then walk forward until we hit the end
# of ADR 4 content (the "Diff-based versioning..." and "Event sourcing..." paragraphs)
# followed by an empty paragraph.

adr4_tbl = None
for child in list(body):
    tag = child.tag.split("}")[1]
    if tag == "p":
        text = "".join(r.text or "" for r in child.findall(".//" + qn("w:t")))
        if "ADR 4" in text and "Store Full Document" in text:
            adr4_tbl = child.getnext()
            break

if adr4_tbl is None:
    print("  WARNING: Could not find ADR 4 table!")
else:
    # Walk forward from the ADR 4 table to find the last non-empty paragraph
    # (which is "Event sourcing: rejected..." or "Diff-based versioning...")
    # followed by empty paragraphs before sectPr
    # We'll insert after the LAST empty paragraph before sectPr (i.e., right before sectPr)

    # Find last meaningful content before sectPr
    all_children = list(body)
    last_content = None
    for child in reversed(all_children):
        tag = child.tag.split("}")[1]
        if tag == "sectPr":
            continue
        if tag == "p":
            text = "".join(r.text or "" for r in child.findall(".//" + qn("w:t")))
            if text.strip():
                last_content = child
                break
        break

    # Find the last element before sectPr (excluding sectPr itself)
    sectPr = body.find(qn("w:sectPr"))
    insert_point = sectPr.getprevious()

    # Build ADR 5 content
    p_heading = make_paragraph_xml("ADR 5 — Use LM Studio for AI Inference")

    adr5_table_rows = [
        ["Field", "Details"],
        ["Title", "ADR 5 — Use LM Studio for Local AI Inference"],
        ["Status", "Accepted"],
    ]
    tbl_adr5 = make_table_xml(adr5_table_rows, col_widths=[1800, 6300])

    p_empty = make_paragraph_xml("")

    p_context_label = make_paragraph_xml("Context")
    p_context = make_paragraph_xml(
        "The AI Writing Assistant feature requires a language model to process text. "
        "The system needs to call an LLM for rewrite, summarize, translate, and restructure "
        "operations without incurring per-request cloud API costs or sending user content "
        "to external servers."
    )
    p_empty2 = make_paragraph_xml("")

    p_decision_label = make_paragraph_xml("Decision")
    p_decision = make_paragraph_xml(
        "We chose LM Studio as the local LLM host. LM Studio runs locally and exposes an "
        "OpenAI-compatible REST API at http://127.0.0.1:1234/v1. The backend uses the openai "
        "Python package as the HTTP client, pointing its base_url at the local LM Studio server. "
        "This means the same client code works with any OpenAI-compatible provider with only "
        "an environment variable change."
    )
    p_empty3 = make_paragraph_xml("")

    p_cons_pos_label = make_paragraph_xml("Consequences — Positive")
    cons_pos = [
        "No API costs — the model runs entirely on local hardware",
        "User content never leaves the local network, preserving privacy",
        "OpenAI-compatible API means the client code is provider-agnostic",
        "Easy to swap the underlying model by changing LM_STUDIO_MODEL",
    ]
    pos_paras = [make_paragraph_xml(c) for c in cons_pos]

    p_cons_neg_label = make_paragraph_xml("Consequences — Negative")
    cons_neg = [
        "LM Studio must be running locally; AI features are unavailable if it is not started",
        "Model quality depends on the model loaded in LM Studio",
        "Local hardware must be capable of running the chosen model",
    ]
    neg_paras = [make_paragraph_xml(c) for c in cons_neg]

    p_empty4 = make_paragraph_xml("")
    p_alt_label = make_paragraph_xml("Alternatives Considered")
    p_alt1 = make_paragraph_xml(
        "OpenAI API: rejected — incurs per-request costs and sends user content to external servers."
    )
    p_alt2 = make_paragraph_xml(
        "Anthropic Claude API: rejected — same concerns as OpenAI API regarding cost and data privacy."
    )
    p_alt3 = make_paragraph_xml(
        "Ollama: considered — similar local LLM hosting with OpenAI-compatible API. "
        "LM Studio chosen for its GUI model management, making it easier to switch models during development."
    )
    p_empty5 = make_paragraph_xml("")

    # Insert all before sectPr
    els_to_insert = [
        make_paragraph_xml(""),
        p_heading,
        tbl_adr5,
        p_empty,
        p_context_label,
        p_context,
        p_empty2,
        p_decision_label,
        p_decision,
        p_empty3,
        p_cons_pos_label,
        *pos_paras,
        p_cons_neg_label,
        *neg_paras,
        p_empty4,
        p_alt_label,
        p_alt1,
        p_alt2,
        p_alt3,
        p_empty5,
    ]

    for el in reversed(els_to_insert):
        sectPr.addprevious(el)

    print("  Done.")


# ═══════════════════════════════════════════════════════════════════════════
# STEP 9 — Update "four ADRs" to "five ADRs" in the intro sentence
# ═══════════════════════════════════════════════════════════════════════════
print("Step 9: Updating ADR count in intro sentence...")

for child in list(body):
    tag = child.tag.split("}")[1]
    if tag == "p":
        text = "".join(r.text or "" for r in child.findall(".//" + qn("w:t")))
        if "following four ADRs" in text:
            for r_el in child.findall(qn("w:r")):
                child.remove(r_el)
            new_r = OxmlElement("w:r")
            new_t = OxmlElement("w:t")
            new_t.text = text.replace("following four ADRs", "following five ADRs")
            new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            new_r.append(new_t)
            child.append(new_r)
            print("  Done.")
            break


# ─── Save ────────────────────────────────────────────────────────────────────
doc.save(DOC_PATH)
print("\nSaved:", DOC_PATH)

# Quick verification
doc2 = Document(DOC_PATH)
print(f"Tables: {len(doc2.tables)}, Paragraphs: {len(doc2.paragraphs)}")
